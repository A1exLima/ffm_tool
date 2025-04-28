import os
import streamlit as st
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from PIL import Image
from io import BytesIO
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF

# === Sessão HTTP com cabeçalhos de navegador e retries ===
session = requests.Session()
session.headers.update({
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/114.0.0.0 Safari/537.36"
    )
})
retries = Retry(
    total=3,
    backoff_factor=1,
    status_forcelist=[429, 500, 502, 503, 504],
    allowed_methods=["GET"],
)
session.mount("https://", HTTPAdapter(max_retries=retries))

# === Caminhos e configuração de página ===
BASE_DIR = os.path.dirname(__file__)
TEMPLATE_RELACIONAMENTO = os.path.join(BASE_DIR, "planilha_modelo_relacionamento.xlsx")
TEMPLATE_AJUDA_CUSTO   = os.path.join(BASE_DIR, "planilha_modelo_ajuda_de_custo.xlsx")

st.set_page_config(page_title="SPOT - Automação FFM", layout="wide")

# === Funções auxiliares ===
def inserir_logo():
    st.image(os.path.join(BASE_DIR, "logo_spot.png"), width=250)

def download_button(label, path, filename):
    if not os.path.isfile(path):
        st.error(f"❌ Template não encontrado: {filename}")
        return
    with open(path, "rb") as f:
        data = f.read()
    st.download_button(
        label=label,
        data=data,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

def ajustar_altura_doc_paragrafo(par):
    pPr = par._element.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepLines'))
    pPr.append(OxmlElement('w:keepNext'))

def inserir_imagem_redimensionada(par, img, largura_max=5.5, altura_max=7):
    bio = BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    w, h = img.size
    scale = min((largura_max*96)/w, (altura_max*96)/h)*1.1
    run = par.add_run()
    run.add_picture(bio, width=Inches(w*scale/96))

def aplicar_fonte_arial(run):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)

def extrair_links_por_relatorio(file):
    wb = load_workbook(file, data_only=True)
    ws = wb.active
    categorias = {
        "FOTO DA AÇÃO E CONQUISTA": list(range(1,5)),
        "FOTO DA LISTA DE PRESENÇA": list(range(5,9)),
        "UMA FOTO DA NOTA FISCAL":   list(range(9,13)),
    }
    resultados = []
    for row in ws.iter_rows(min_row=2):
        num = row[0].value
        if not num: continue
        grupos = {}
        for cat, idxs in categorias.items():
            urls = []
            for i in idxs:
                if i < len(row):
                    cell = row[i]
                    url = cell.hyperlink.target if cell.hyperlink else None
                    if not url and isinstance(cell.value, str) and cell.value.startswith(("http://","https://")):
                        url = cell.value
                    if url:
                        urls.append(url)
            grupos[cat] = urls
        resultados.append((num, grupos))
    return resultados

def extrair_links_por_ajuda_custo(file):
    wb = load_workbook(file, data_only=True)
    ws = wb.active
    header = [c.value for c in ws[1]]
    cat1 = header[1] or "Comprovante Capturado"
    cat2 = header[2] or "Outras Evidências"
    resultados = []
    for row in ws.iter_rows(min_row=2):
        idr = row[0].value
        if not idr: continue
        grupos = {
            cat1: [str(row[1].value)] if len(row)>1 and row[1].value else [],
            cat2: [str(row[2].value)] if len(row)>2 and row[2].value else []
        }
        resultados.append((idr, grupos))
    return resultados

def pdf_para_imagens(pdf_bytes):
    imgs = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc_pdf:
        for page in doc_pdf:
            pix = page.get_pixmap(dpi=150)
            imgs.append(Image.open(BytesIO(pix.tobytes("png"))))
    return imgs

# === Interface ===
inserir_logo()
opcao = st.sidebar.radio("Opções", ["Ação de Relacionamento", "Ajuda de Custo"])

# --- Ação de Relacionamento ---
if opcao == "Ação de Relacionamento":
    st.title("Automação FFM - Ação de Relacionamento")
    download_button(
        "Planilha Modelo",
        TEMPLATE_RELACIONAMENTO,
        os.path.basename(TEMPLATE_RELACIONAMENTO)
    )
    uploaded = st.file_uploader(
        "📂 Envie a planilha de Relacionamento (.xlsx)",
        type="xlsx",
        key="rel_upload"
    )
    if uploaded:
        info = extrair_links_por_relatorio(uploaded)
        total = len(info)
        if not info:
            st.error("❌ Nenhum link encontrado.")
        else:
            st.success(f"✅ {total} relatórios encontrados.")
            if st.button("📝 Gerar Documento Word", key="btn_rel"):
                doc = Document()
                log = st.empty()
                for i, (num_rel, grupos) in enumerate(info, 1):
                    with st.spinner("Processando"):
                        log.markdown(
                            f"🔄 <span style='color:red'>Linha {i}/{total}</span> - "
                            f"Processando Relatório <span style='color:red'>{num_rel}</span>",
                            unsafe_allow_html=True
                        )
                        for categoria, links in grupos.items():
                            if not links:
                                continue
                            doc.add_page_break()
                            p = doc.add_paragraph()
                            ajustar_altura_doc_paragrafo(p)
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = p.add_run(f"Relatório: {num_rel} — {categoria}")
                            aplicar_fonte_arial(run)
                            for link in links:
                                if not link.startswith("http"):
                                    link = "https://" + link.strip()
                                success = False
                                for attempt in range(1, 4):
                                    try:
                                        resp = session.get(link, timeout=(5,60), stream=True)
                                        resp.raise_for_status()
                                        ct = resp.headers.get("Content-Type","")
                                        imgs = pdf_para_imagens(resp.content) if "pdf" in ct else [Image.open(BytesIO(resp.content)).convert("RGB")]
                                        for img in imgs:
                                            p_img = doc.add_paragraph()
                                            p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                            inserir_imagem_redimensionada(p_img, img)
                                        success = True
                                        break
                                    except Exception as e:
                                        if attempt == 3:
                                            pe = doc.add_paragraph()
                                            pe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                            re = pe.add_run(f"⚠️ Erro após 3 tentativas: {e}")
                                            aplicar_fonte_arial(re)
                                if not success:
                                    st.warning(f"Falha ao carregar: {link}")
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                log.empty()
                st.success("✅ Documento gerado!")
                st.download_button(
                    "📥 Baixar Word - Relacionamento",
                    data=buf,
                    file_name="evidencias_acao_relacionamento.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# --- Ajuda de Custo ---
elif opcao == "Ajuda de Custo":
    st.title("Automação FFM - Ajuda de Custo")
    download_button(
        "Planilha Modelo",
        TEMPLATE_AJUDA_CUSTO,
        os.path.basename(TEMPLATE_AJUDA_CUSTO)
    )
    uploaded = st.file_uploader(
        "📂 Envie a planilha de Ajuda de Custo (.xlsx)",
        type="xlsx",
        key="ajc_upload"
    )
    if uploaded:
        info = extrair_links_por_ajuda_custo(uploaded)
        total = len(info)
        if not info:
            st.error("❌ Nenhum link encontrado.")
        else:
            st.success(f"✅ {total} reembolsos encontrados.")
            if st.button("📝 Gerar Documento Word - Ajuda de Custo", key="btn_ajc"):
                doc = Document()
                log = st.empty()
                for i, (idr, grupos) in enumerate(info, 1):
                    with st.spinner("Processando"):
                        log.markdown(
                            f"🔄 <span style='color:red'>Linha {i}/{total}</span> - "
                            f"Processando Id do reembolso <span style='color:red'>{idr}</span>",
                            unsafe_allow_html=True
                        )
                        for categoria, links in grupos.items():
                            if not links:
                                continue
                            doc.add_page_break()
                            p = doc.add_paragraph()
                            ajustar_altura_doc_paragrafo(p)
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = p.add_run(f"Reembolso: {idr} — {categoria}")
                            aplicar_fonte_arial(run)
                            for link in links:
                                if not link.startswith("http"):
                                    link = "https://" + link.strip()
                                success = False
                                for attempt in range(1, 4):
                                    try:
                                        resp = session.get(link, timeout=(5,60), stream=True)
                                        resp.raise_for_status()
                                        ct = resp.headers.get("Content-Type","")
                                        imgs = pdf_para_imagens(resp.content) if "pdf" in ct else [Image.open(BytesIO(resp.content)).convert("RGB")]
                                        for img in imgs:
                                            p_img = doc.add_paragraph()
                                            p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                            inserir_imagem_redimensionada(p_img, img)
                                        success = True
                                        break
                                    except Exception as e:
                                        if attempt == 3:
                                            pe = doc.add_paragraph()
                                            pe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                            re = pe.add_run(f"⚠️ Erro após 3 tentativas: {e}")
                                            aplicar_fonte_arial(re)
                                if not success:
                                    st.warning(f"Falha ao carregar: {link}")
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                log.empty()
                st.success("✅ Documento gerado!")
                st.download_button(
                    "📥 Baixar Word - Ajuda de Custo",
                    data=buf,
                    file_name="evidencias_ajuda_de_custo.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
